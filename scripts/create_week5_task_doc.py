from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_FILE = OUT_DIR / "Week_5_Multi_Agent_WORD_READY.docx"
RTF_FILE = OUT_DIR / "Week_5_Multi_Agent_WORD_READY.rtf"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(document: Document, value: str) -> None:
    paragraph = document.add_heading(value, level=1)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor(31, 78, 121)


def subheading(document: Document, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(value)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)


def body(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(value)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08


def bullets(document: Document, values: list[str]) -> None:
    for value in values:
        paragraph = document.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)


def code(document: Document, value: str) -> None:
    for line in value.splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
    document.add_paragraph()


def table(document: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        write_cell(tbl.rows[0].cells[i], header, True)
        shade_cell(tbl.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value)
    document.add_paragraph()


def escape_rtf(value: str) -> str:
    return value.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")


def write_rtf(docx_path: Path, rtf_path: Path) -> None:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    paragraphs = []
    for paragraph in root.findall(".//w:p", namespace):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespace)).strip()
        if text:
            paragraphs.append(text)
    lines = [r"{\rtf1\ansi\deff0", r"{\fonttbl{\f0 Calibri;}{\f1 Consolas;}}"]
    for text in paragraphs:
        font = r"\f1\fs18" if text.startswith(("|", "Voltstream_application/")) else r"\f0\fs22"
        bold = r"\b " if text[:2].isdigit() or text.startswith("Week 5") else ""
        bold_end = r"\b0" if bold else ""
        lines.append(rf"\pard{font} {bold}{escape_rtf(text)}{bold_end}\par")
    lines.append("}")
    rtf_path.write_text("\n".join(lines), encoding="utf-8")


def add_file_structure(document: Document) -> None:
    body(document, "Recommended submission folder structure based on the actual Week 5 files:")
    code(
        document,
        "Voltstream_application/\n"
        "|\n"
        "|-- README.md\n"
        "|\n"
        "|-- .env.example\n"
        "|\n"
        "|-- backend/\n"
        "|   |\n"
        "|   |-- main.py\n"
        "|   |\n"
        "|   |-- api.py\n"
        "|   |\n"
        "|   |-- data_models.py\n"
        "|   |\n"
        "|   |-- database.py\n"
        "|   |\n"
        "|   `-- agents/\n"
        "|       |\n"
        "|       |-- __init__.py\n"
        "|       |\n"
        "|       |-- comments.py\n"
        "|       |\n"
        "|       |-- analyst_agent.py\n"
        "|       |\n"
        "|       |-- advisor_agent.py\n"
        "|       |\n"
        "|       |-- orchestrator_agent.py\n"
        "|       |\n"
        "|       `-- runner_agent.py\n"
        "|\n"
        "|-- frontend/\n"
        "|   |\n"
        "|   `-- src/\n"
        "|       |\n"
        "|       |-- api.js\n"
        "|       |\n"
        "|       `-- components/DeviceAgent.jsx\n"
        "|\n"
        "`-- docs/\n"
        "    |\n"
        "    `-- Week_5_Multi_Agent_WORD_READY.docx",
    )


def add_architecture(document: Document) -> None:
    heading(document, "11. Architecture of Multi-Agents Working Flow")
    body(
        document,
        "The Week 5 architecture is a layered multi-agent workflow. The user interacts with the frontend or API client. "
        "FastAPI receives the prompt and calls runner_agent.py. The runner does not perform all work by itself; it "
        "coordinates the Orchestrator, Analyst, and Advisor sequence. Each agent has a focused responsibility, and the "
        "final response is built from their combined output.",
    )
    subheading(document, "Architecture flowchart")
    code(
        document,
        "Frontend / API Client\n"
        "   |\n"
        "   |  POST /api/v1/agent\n"
        "   v\n"
        "FastAPI Endpoint - backend/api.py\n"
        "   |\n"
        "   |  request.agent_type = energy\n"
        "   v\n"
        "Runner Layer - backend/agents/runner_agent.py\n"
        "   |\n"
        "   |  capture_tool_trace()\n"
        "   v\n"
        "Orchestrator Agent - orchestrator_agent.py\n"
        "   |\n"
        "   |-- route_energy_usage_request\n"
        "   v\n"
        "Analyst Agent - analyst_agent.py\n"
        "   |\n"
        "   |-- reads usage history\n"
        "   |-- reads active devices\n"
        "   |-- reads billing summary\n"
        "   |-- calculates grid / solar / net usage\n"
        "   |-- finds top device loads\n"
        "   v\n"
        "Structured Analysis Result\n"
        "   |\n"
        "   |-- route_analysis_to_advisor\n"
        "   v\n"
        "Advisor Agent - advisor_agent.py\n"
        "   |\n"
        "   |-- reviews top loads\n"
        "   |-- creates recommendations\n"
        "   |-- calculates target savings\n"
        "   |-- checks budget risk\n"
        "   v\n"
        "Recommendation Result\n"
        "   |\n"
        "   v\n"
        "Orchestrator Agent - final formatting\n"
        "   |\n"
        "   |-- combines usage + advice + billing\n"
        "   v\n"
        "AgentResponse\n"
        "   |\n"
        "   |-- answer\n"
        "   |-- tool_calls\n"
        "   `-- ENERGY_AGENT_LOOP",
    )
    subheading(document, "Architecture explanation")
    bullets(
        document,
        [
            "Frontend/API Client is the user-facing layer. It sends the prompt to the backend.",
            "FastAPI endpoint validates the request and decides which runner function should handle it.",
            "Runner Agent is the execution layer. It starts trace capture and calls each specialist step in order.",
            "Orchestrator Agent is the routing and final-formatting layer. It decides the flow and prepares the final response.",
            "Analyst Agent is the data-understanding layer. It reads usage, solar, device, and billing data.",
            "Advisor Agent is the recommendation layer. It converts analysis into actions the user can follow.",
            "comments.py supports the whole architecture by storing names, prompts, and ENERGY_AGENT_LOOP.",
            "AgentResponse is the output layer. It returns answer, tool_calls, and agent_loop to the frontend.",
        ],
    )


def add_learning_objectives(document: Document) -> None:
    heading(document, "14. Learning Objectives of Week 5 Task")
    body(
        document,
        "The learning objectives of Week 5 focus on understanding how a backend agent system can grow from one agent "
        "into a coordinated multi-agent workflow. The main learning is how to separate analysis, advice, routing, and "
        "execution into different files while still returning one clean answer to the user.",
    )
    bullets(
        document,
        [
            "Understand how multi-agent systems split work between specialist agents.",
            "Understand the role of the Orchestrator Agent in routing, sequencing, and final answer formatting.",
            "Understand the role of the Analyst Agent in reading usage history, solar offset, device loads, and billing data.",
            "Understand the role of the Advisor Agent in creating practical recommendations and target savings.",
            "Understand how comments.py centralizes model name, agent names, instructions, and the visible agent loop.",
            "Understand how runner_agent.py executes the complete flow and records tool_calls.",
            "Understand how FastAPI exposes the multi-agent workflow through /api/v1/agent.",
            "Understand how tool_calls and ENERGY_AGENT_LOOP make the internal workflow explainable during demo.",
        ],
    )


def add_week5_topics(document: Document) -> None:
    heading(document, "15. Week 5 Topics Covered")
    body(
        document,
        "The Week 5 task screenshot focuses on Agentic AI at the intermediate level. The important topics are not only "
        "the project files, but also the concepts behind deploying, invoking, tracing, and explaining an agent workflow. "
        "The topics below summarize the Week 5 learning points in simple project-related language.",
    )
    subheading(document, "AgentCore")
    body(
        document,
        "AgentCore is the managed runtime concept from the Week 5 task. Compared with running an agent only on a local "
        "machine, AgentCore adds a cloud execution layer where the agent can be deployed, invoked, tracked, and managed. "
        "In documentation terms, it explains how the agent moves from local development toward a production-style runtime.",
    )
    subheading(document, "Strands Agent")
    body(
        document,
        "A Strands agent is the agent implementation style referenced in the Week 5 task. The main idea is to build an "
        "agent with clear tools, instructions, and a reasoning flow. In the VoltStream project, this maps to a multi-agent "
        "structure where the Orchestrator, Analyst, and Advisor each have focused responsibilities.",
    )
    subheading(document, "Sessions")
    body(
        document,
        "A session is used to keep track of an invocation context. The Week 5 screenshot asks to show the session ID and "
        "execution logs. This is important because a session ID proves that the request was not just a one-line response; "
        "it went through a tracked agent invocation flow.",
    )
    subheading(document, "Memory")
    body(
        document,
        "Memory means the ability to preserve useful context across agent interactions. In a production agent system, "
        "memory can help the assistant remember previous messages, user preferences, or task state. For VoltStream, memory "
        "could help continue an energy conversation across multiple prompts instead of treating every prompt as isolated.",
    )
    subheading(document, "Lifecycle management")
    body(
        document,
        "Lifecycle management refers to how an agent is packaged, started, invoked, monitored, updated, and stopped. "
        "This matters in Week 5 because deployment is more advanced than local execution. A deployed agent needs a stable "
        "runtime lifecycle so it can be called repeatedly and observed through logs.",
    )
    subheading(document, "Invocation through Lambda")
    body(
        document,
        "Invocation through Lambda means an external request can call a Lambda function, and Lambda can trigger the agent "
        "runtime. The screenshot specifically mentions the flow Lambda -> AgentCore -> Strands agent -> Bedrock -> response. "
        "This shows how an agent can be exposed through a cloud function instead of being manually run from a developer machine.",
    )
    subheading(document, "Execution logs")
    body(
        document,
        "Execution logs are records of what happened during an invocation. They help prove that the agent was called, which "
        "session was used, which route was followed, and whether the response was generated successfully. In the VoltStream "
        "documentation, tool_calls and ENERGY_AGENT_LOOP provide a local project-style trace of this same idea.",
    )
    subheading(document, "Two-agent system: Device Control Agent + Energy Advisor Agent")
    body(
        document,
        "The Week 5 task asks for a two-agent system idea: Device Control Agent plus Energy Advisor Agent. In the current "
        "VoltStream project, this is represented through specialist files. The device-control side handles device state, "
        "while the energy workflow uses Analyst and Advisor logic to understand usage and recommend savings actions.",
    )
    subheading(document, "What AgentCore adds compared with local execution")
    bullets(
        document,
        [
            "A managed runtime instead of only running Python locally.",
            "Session IDs that help track invocations.",
            "Execution logs for debugging and demonstration.",
            "Lifecycle management for deployment and repeated runtime use.",
            "A cloud invocation path through Lambda.",
            "A clearer route toward production deployment.",
        ],
    )


def build_document() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Week 5 - Agentic AI Intermediate")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("VoltStream Multi-Agent Energy Workflow").font.size = Pt(12)

    heading(doc, "1. Title")
    body(doc, "Week 5 - Agentic AI Intermediate: VoltStream Multi-Agent Energy Workflow.")

    heading(doc, "2. Introduction")
    body(doc, "The Week 5 VoltStream task documents the multi-agent energy workflow already present in the backend agents folder. Instead of creating unrelated folders, this document uses the actual project files: comments.py, analyst_agent.py, advisor_agent.py, orchestrator_agent.py, and runner_agent.py.")
    body(doc, "The user sends an energy question to /api/v1/agent. The Orchestrator routes the request to the Analyst. The Analyst reads VoltStream usage and billing data. The Advisor creates recommendations. The Orchestrator formats the final answer.")

    heading(doc, "3. Objective")
    body(doc, "The objective is to explain and demonstrate a multi-agent energy assistant for VoltStream. The assistant should analyze energy usage, identify top loads, include billing context, provide recommendations, and return the full agent loop trace.")

    heading(doc, "4. Use Cases")
    bullets(doc, [
        "Weekly usage summary: user asks for this week's or last week's energy usage.",
        "Daily or monthly usage summary: user asks about today, yesterday, this month, or last month.",
        "Top device load analysis: Analyst estimates which active devices contribute most.",
        "Billing guidance: Analyst includes projected bill, budget limit, and solar credit.",
        "Savings advice: Advisor gives practical steps to reduce grid usage and cost.",
        "Traceable demo: response shows Orchestrator, Analyst, Advisor, and final answer steps.",
    ])

    heading(doc, "5. File Structure")
    add_file_structure(doc)

    heading(doc, "6. Use of Every Week 5 File")
    table(doc, ["File", "Use in Week 5 task"], [
        ("backend/agents/comments.py", "Stores DEFAULT_MODEL, agent names, instructions, ENERGY_AGENT_LOOP, and capture_tool_trace."),
        ("backend/agents/analyst_agent.py", "Analyzes usage period, grid usage, solar offset, net usage, top devices, and billing context."),
        ("backend/agents/advisor_agent.py", "Creates recommendations from analyst output and calculates target savings."),
        ("backend/agents/orchestrator_agent.py", "Routes between Analyst and Advisor, records route steps, and formats the final answer."),
        ("backend/agents/runner_agent.py", "Runs the full multi-agent flow and returns answer, tool_calls, and agent_loop."),
        ("backend/api.py", "Exposes POST /api/v1/agent and returns AgentResponse."),
        ("backend/data_models.py", "Defines AgentRequest, AgentResponse, and AgentToolCall schemas."),
        ("backend/database.py", "Provides usage history, device, and billing data used by Analyst Agent."),
        ("frontend/src/api.js", "Sends frontend requests to backend endpoints."),
        ("frontend/src/components/DeviceAgent.jsx", "Displays prompt input, answer, tool calls, and agent loop."),
    ])

    heading(doc, "7. Working Flow of Week 5 Task")
    body(doc, "The working flow starts from the user prompt and ends with a structured AgentResponse. The system passes through API, runner, orchestrator, analyst, advisor, and final formatting layers.")
    code(doc, "User Energy Prompt\n   |\n   v\nPOST /api/v1/agent\n   |\n   v\nrun_energy_usage_agent(message)\n   |\n   v\nOrchestrator Agent\n   |\n   v\nAnalyst Agent -> analyze_usage(message)\n   |\n   v\nAdvisor Agent -> generate_recommendations(analysis)\n   |\n   v\nOrchestrator Agent -> _format_energy_answer(...)\n   |\n   v\nAgentResponse(answer, tool_calls, ENERGY_AGENT_LOOP)")

    heading(doc, "8. Week 5 Checkpoint")
    bullets(doc, [
        "Send an energy-related request to POST /api/v1/agent.",
        "Show that the request is routed to Analyst Agent.",
        "Show that Analyst Agent reads usage history, active devices, solar offset, and billing context.",
        "Show that the result is routed to Advisor Agent.",
        "Show that Advisor Agent returns recommendations and target savings.",
        "Show the final Orchestrator response and explain the tool_calls trace.",
    ])

    heading(doc, "9. Hands-On Activities and Deliverables")
    bullets(doc, [
        "Created specialist logic for Analyst Agent, Advisor Agent, and Orchestrator Agent.",
        "Centralized agent names, prompts, and loop descriptions in comments.py.",
        "Connected the flow through run_energy_usage_agent in runner_agent.py.",
        "Returned answer, tool_calls, and ENERGY_AGENT_LOOP from the backend API.",
        "Prepared documentation with file structure, file usage, working flow, endpoints, and code logic.",
    ])

    heading(doc, "10. Main Multi-Agent Loop")
    bullets(doc, [
        "User prompt received by FastAPI /api/v1/agent.",
        "Orchestrator Agent routes request to Analyst Agent.",
        "Analyst Agent reads usage period, usage history, devices, solar data, and billing context.",
        "Orchestrator Agent sends analysis to Advisor Agent.",
        "Advisor Agent converts analysis into recommendations.",
        "Orchestrator Agent formats and returns the final answer.",
    ])

    add_architecture(doc)

    heading(doc, "12. API Endpoints Related to Week 5")
    table(doc, ["Method", "Endpoint", "Purpose"], [
        ("POST", "/api/v1/agent", "Main Week 5 endpoint for the multi-agent energy workflow."),
        ("GET", "/api/v1/analytics/history?period=daily|weekly|monthly", "Provides usage history used by Analyst Agent."),
        ("GET", "/api/v1/devices", "Provides device list and active loads."),
        ("GET", "/api/v1/billing/summary", "Provides billing context."),
        ("GET", "/api/v1/billing/trend", "Provides billing trend context."),
        ("GET", "/health", "Checks backend health before demo."),
    ])
    code(doc, '{\n  "message": "Show this week\'s energy usage and give recommendations",\n  "agent_type": "energy",\n  "user_id": "voltstream-user",\n  "session_id": "voltstream-session"\n}')

    heading(doc, "13. Main Code Logic")
    subheading(doc, "12.1 comments.py")
    code(doc, 'DEFAULT_MODEL = os.getenv("ADK_MODEL", "gemini-2.5-flash")\nORCHESTRATOR_AGENT_NAME = "voltstream_orchestrator_agent"\nANALYST_AGENT_NAME = "voltstream_analyst_agent"\nADVISOR_AGENT_NAME = "voltstream_advisor_agent"\nENERGY_AGENT_LOOP = [...]')
    body(doc, "comments.py keeps shared names, prompts, and loop text in one place. This keeps all agents consistent.")
    subheading(doc, "12.2 analyst_agent.py")
    code(doc, "def analyze_usage(message: str) -> dict[str, Any]:\n    period = _usage_period(message)\n    usage_history = get_usage_history(period)\n    devices = get_devices()\n    billing = get_billing_summary()\n    return analysis")
    body(doc, "Analyst Agent converts VoltStream usage and billing data into structured facts.")
    subheading(doc, "12.3 advisor_agent.py")
    code(doc, "def generate_recommendations(analysis: dict[str, Any]) -> dict[str, Any]:\n    recommendations = []\n    target_savings_kwh = round(total_grid_kwh * 0.12, 1)\n    return {\"recommendations\": recommendations, \"target_savings_kwh\": target_savings_kwh}")
    body(doc, "Advisor Agent converts the analysis into action-oriented recommendations.")
    subheading(doc, "12.4 orchestrator_agent.py")
    code(doc, "def _format_energy_answer(message, analysis, advice):\n    lines = [grid_usage, solar_offset, net_usage]\n    lines.extend(recommendations)\n    return \"\\n\".join(lines)")
    body(doc, "Orchestrator Agent controls routing and final response format.")
    subheading(doc, "12.5 runner_agent.py")
    code(doc, "def run_energy_usage_agent(message: str) -> dict[str, Any]:\n    _orchestrator_agent_route(..., ANALYST_AGENT_NAME)\n    analysis = analyze_usage(message)\n    _orchestrator_agent_route(..., ADVISOR_AGENT_NAME)\n    advice = generate_recommendations(analysis)\n    answer = _format_energy_answer(message, analysis, advice)\n    return {\"answer\": answer, \"tool_calls\": trace, \"agent_loop\": ENERGY_AGENT_LOOP}")
    body(doc, "runner_agent.py executes the full Week 5 flow in the correct order.")

    add_learning_objectives(doc)
    add_week5_topics(doc)

    heading(doc, "16. Demo Request and Expected Response")
    body(doc, "Expected response: Orchestrator routes to Analyst, Analyst returns analysis, Orchestrator routes to Advisor, Advisor returns recommendations, and Orchestrator returns the final answer. The API response includes answer, tool_calls, and ENERGY_AGENT_LOOP.")

    heading(doc, "17. Deliverables Completed")
    bullets(doc, [
        "Analyst Agent documented.",
        "Advisor Agent documented.",
        "Orchestrator Agent documented.",
        "Runner Agent documented.",
        "comments.py shared instructions and trace logic documented.",
        "API endpoints, working flow, file structure, and code logic documented.",
    ])

    doc.save(DOCX_FILE)
    write_rtf(DOCX_FILE, RTF_FILE)


if __name__ == "__main__":
    build_document()
    print(DOCX_FILE)
